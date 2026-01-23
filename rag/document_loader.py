"""
문서 로더 v6.2 - section_path 계층 추적 추가
- PDF, DOCX, HTML, 이미지 등 다양한 형식 지원
- 표(Table) 파싱 지원
- section_path: "5 > 5.1 > 5.1.1" 형태의 계층 경로
- section_path_readable: "5 절차 > 5.1 문서체계 > 5.1.1 Level 1" 형태
"""

import re
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field
import io
import tempfile
import os


@dataclass
class ContentBlock:
    """문서의 의미 단위 블록"""
    text: str
    block_type: str  # title, paragraph, table, list, article 등
    level: int = 0
    page: Optional[int] = None
    section: Optional[str] = None
    metadata: Dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """파싱된 문서"""
    text: str
    blocks: List[ContentBlock]
    metadata: Dict
    tables: List[Dict] = field(default_factory=list)  # 표 데이터


def get_supported_extensions() -> list:
    """지원되는 파일 확장자"""
    return [".txt", ".md", ".pdf", ".docx", ".doc", ".html", ".htm", ".csv", ".xlsx", ".pptx", ".png", ".jpg", ".jpeg"]


def load_document(filename: str, content: bytes) -> ParsedDocument:
    """
    문서 로드 및 파싱 (메인 진입점)
    
    Returns:
        ParsedDocument: 파싱된 문서 (text + blocks + metadata + tables)
    """
    ext = Path(filename).suffix.lower()

    # DOCX: 하이브리드 방식 (Docling 표 + python-docx 텍스트)
    if ext in [".docx", ".doc"]:
        return _load_docx_hybrid(filename, content)

    # PDF, PPTX, XLSX, HTML, 이미지 → Docling 시도
    if ext in [".pdf", ".pptx", ".xlsx", ".html", ".htm", ".png", ".jpg", ".jpeg"]:
        try:
            return _load_with_docling(filename, content, ext)
        except ImportError:
            print("⚠️ Docling not installed, falling back to basic parser")
            if ext == ".pdf":
                return _load_pdf_basic(filename, content)
            elif ext in [".html", ".htm"]:
                return _load_html_basic(filename, content)
        except Exception as e:
            print(f"⚠️ Docling failed: {e}, falling back to basic parser")
            if ext == ".pdf":
                return _load_pdf_basic(filename, content)
            elif ext in [".html", ".htm"]:
                return _load_html_basic(filename, content)

    # 텍스트 기반 파일
    if ext in [".txt", ".md"]:
        text = _decode_text(content)
        if _is_article_document(text):
            return _parse_articles(text, filename, ext)
        return _parse_plain_text(text, filename, ext)

    if ext == ".csv":
        return _load_csv(filename, content)

    # 기본 텍스트 처리
    text = _decode_text(content)
    return _parse_plain_text(text, filename, ext)


def _load_docx_hybrid(filename: str, content: bytes) -> ParsedDocument:
    """
    DOCX 하이브리드 파싱
    - 문서 순서대로 텍스트 추출 (중요!)
    - 표(Table)를 가독성 좋은 형태로 변환
    """
    tables_data = []
    
    # python-docx로 문서 순서대로 파싱
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        return ParsedDocument(
            text="python-docx가 설치되지 않았습니다.",
            blocks=[],
            metadata={"file_name": filename, "error": "python-docx not installed"},
            tables=[]
        )
    
    doc = Document(io.BytesIO(content))
    all_text = []
    current_section = None  # 현재 섹션 추적
    
    # 문서 순서대로 순회 (핵심!)
    for element in doc.element.body:
        # 단락(Paragraph)
        if element.tag.endswith('p'):
            para = Paragraph(element, doc)
            text = para.text.strip()
            if text:
                all_text.append(text)
                # 섹션 감지 (예: "3. 책임 및 역할")
                section_match = re.match(r'^(\d+(?:\.\d+)?)\.\s+(.+)', text)
                if section_match:
                    current_section = section_match.group(2).strip()
        
        # 표(Table)
        elif element.tag.endswith('tbl'):
            table = Table(element, doc)
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            
            if rows:
                # 표를 가독성 좋은 텍스트로 변환
                table_text = _format_table_as_text(rows, current_section)
                all_text.append(table_text)
                tables_data.append({"rows": rows, "source": "python-docx"})
    
    full_text = '\n'.join(all_text)
    
    # 조항 단위 블록 생성 (section_path 포함)
    blocks = _extract_article_blocks(full_text)
    
    # 메타데이터
    metadata = {
        "file_name": filename,
        "file_type": "docx",
        "title": _extract_title(full_text),
        "table_count": len(tables_data),
        "parser": "python-docx (ordered)"
    }
    metadata.update(_extract_sop_metadata(full_text))
    
    return ParsedDocument(
        text=full_text,
        blocks=blocks,
        metadata=metadata,
        tables=tables_data
    )


def _load_docx_basic(filename: str, content: bytes) -> ParsedDocument:
    """기본 DOCX 파싱"""
    try:
        from docx import Document
    except ImportError:
        return ParsedDocument(
            text="DOCX 파싱 라이브러리(python-docx)가 설치되지 않았습니다.",
            blocks=[],
            metadata={"file_name": filename, "error": "python-docx not installed"}
        )

    doc = Document(io.BytesIO(content))
    blocks = []
    all_text = []
    tables = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 스타일로 타입 결정
            style_name = para.style.name.lower() if para.style else ""
            block_type = "title" if "heading" in style_name else "paragraph"
            blocks.append(ContentBlock(text=text, block_type=block_type))
            all_text.append(text)

    # 표 추출
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
            all_text.append(' | '.join(cells))
        if rows:
            tables.append({"rows": rows})
            blocks.append(ContentBlock(
                text='\n'.join([' | '.join(r) for r in rows]),
                block_type="table"
            ))

    full_text = '\n\n'.join(all_text)

    if _is_article_document(full_text):
        article_blocks = _extract_article_blocks(full_text)
        if article_blocks:
            blocks = article_blocks

    metadata = {
        "file_name": filename,
        "file_type": "docx",
        "title": _extract_title(full_text),
        "table_count": len(tables),
        "parser": "python-docx"
    }
    metadata.update(_extract_sop_metadata(full_text))

    return ParsedDocument(text=full_text, blocks=blocks, metadata=metadata, tables=tables)


def _format_table_as_text(rows: List[List[str]], section_title: str = None) -> str:
    """
    표를 가독성 좋은 텍스트로 변환
    - 2열 표: 키-값 형태
    - 다열 표: 헤더 + 행 형태
    """
    if not rows:
        return ""
    
    # 열 개수 확인
    num_cols = max(len(row) for row in rows)
    
    # 2열 표: 키-값 형태
    if num_cols == 2:
        lines = []
        if section_title:
            lines.append(f"[표: {section_title}]")
        
        # 첫 행이 헤더인지 확인 (둘 다 짧은 텍스트면 헤더로 간주)
        first_row = rows[0] if rows else []
        is_header = len(first_row) >= 2 and len(first_row[0]) < 10 and len(first_row[1]) < 10
        
        data_rows = rows[1:] if is_header else rows
        
        for row in data_rows:
            if len(row) >= 2:
                key = row[0].strip()
                value = row[1].strip()
                if key and value:
                    lines.append(f"• {key}: {value}")
                elif key:
                    lines.append(f"• {key}")
        
        return '\n'.join(lines)
    
    # 다열 표: 헤더 + 행 형태
    else:
        lines = []
        if section_title:
            lines.append(f"[표: {section_title}]")
        
        # 첫 행을 헤더로 가정
        if rows:
            header = rows[0]
            for row in rows[1:]:
                row_parts = []
                for i, cell in enumerate(row):
                    if cell.strip():
                        col_name = header[i] if i < len(header) else f"열{i+1}"
                        row_parts.append(f"{col_name}: {cell.strip()}")
                if row_parts:
                    lines.append("• " + " | ".join(row_parts))
        
        return '\n'.join(lines)


# ═══════════════════════════════════════════════════════════════════════════
# Docling 기반 파서 (핵심!)
# ═══════════════════════════════════════════════════════════════════════════

def _load_with_docling(filename: str, content: bytes, ext: str) -> ParsedDocument:
    """Docling을 사용한 고급 문서 파싱"""
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions

    # 임시 파일 생성
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Docling 컨버터 설정
        converter = DocumentConverter()

        # 문서 변환
        result = converter.convert(tmp_path)
        doc = result.document

        # 전체 텍스트
        full_text = doc.export_to_markdown()

        # 블록 추출
        blocks = []
        tables = []

        for item in doc.iterate_items():
            element = item[1] if isinstance(item, tuple) else item

            # 텍스트 추출
            if hasattr(element, 'text'):
                text = element.text
            elif hasattr(element, 'export_to_markdown'):
                text = element.export_to_markdown()
            else:
                continue

            if not text or not text.strip():
                continue

            # 블록 타입 결정
            block_type = "paragraph"
            element_type = type(element).__name__.lower()

            if "title" in element_type or "heading" in element_type:
                block_type = "title"
            elif "table" in element_type:
                block_type = "table"
                # 표 데이터 추출
                table_data = _extract_table_data(element)
                if table_data:
                    tables.append(table_data)
            elif "list" in element_type:
                block_type = "list"

            # 페이지 번호
            page_num = None
            if hasattr(element, 'prov') and element.prov:
                for prov in element.prov:
                    if hasattr(prov, 'page_no'):
                        page_num = prov.page_no
                        break

            blocks.append(ContentBlock(
                text=text.strip(),
                block_type=block_type,
                page=page_num,
                metadata={"source": "docling"}
            ))

        # 조항 패턴 감지 및 재파싱
        if _is_article_document(full_text):
            article_blocks = _extract_article_blocks(full_text)
            if article_blocks:
                blocks = article_blocks

        # 메타데이터 추출
        metadata = {
            "file_name": filename,
            "file_type": ext,
            "title": _extract_title(full_text),
            "total_pages": _count_pages(doc),
            "table_count": len(tables),
            "parser": "docling"
        }

        # SOP 메타데이터 추출
        sop_meta = _extract_sop_metadata(full_text)
        metadata.update(sop_meta)

        return ParsedDocument(
            text=full_text,
            blocks=blocks,
            metadata=metadata,
            tables=tables
        )

    finally:
        # 임시 파일 삭제
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _extract_table_data(element) -> Optional[Dict]:
    """표 요소에서 데이터 추출"""
    try:
        if hasattr(element, 'export_to_dataframe'):
            df = element.export_to_dataframe()
            return {
                "headers": list(df.columns),
                "rows": df.values.tolist(),
                "markdown": element.export_to_markdown() if hasattr(element, 'export_to_markdown') else str(df)
            }
        elif hasattr(element, 'data'):
            data = element.data
            if hasattr(data, 'grid'):
                return {
                    "grid": data.grid,
                    "markdown": element.export_to_markdown() if hasattr(element, 'export_to_markdown') else ""
                }
    except Exception as e:
        print(f"표 추출 실패: {e}")
    return None


def _count_pages(doc) -> int:
    """문서 페이지 수"""
    try:
        if hasattr(doc, 'pages'):
            return len(doc.pages)
    except Exception:
        pass
    return 0


# ═══════════════════════════════════════════════════════════════════════════
# 조항 파싱 (SOP/법률 문서) - section_path 계층 추적 추가
# ═══════════════════════════════════════════════════════════════════════════

ARTICLE_PATTERNS = [
    # 한글 조항
    (r'^제\s*(\d+)\s*조\s*(.*)', 'article'),
    (r'^제\s*(\d+)\s*장\s*(.*)', 'chapter'),
    (r'^제\s*(\d+)\s*절\s*(.*)', 'section'),
    
    # 🔥 "제 N레벨" 형식 (이 문서 전용)
    (r'^제\s*(\d+)\s*레벨\s*[:\(]?\s*(.+)', 'level'),  # "제 1레벨(품질매뉴얼):"
    
    # 숫자형 (점 있음): 구체적인 것 먼저!
    (r'^(\d+\.\d+\.\d+)\s+([가-힣A-Za-z].+)', 'subsubsection'),  # "5.1.1 Level 1"
    (r'^(\d+\.\d+)\s+([가-힣A-Za-z].+)', 'subsection'),          # "6.1 사전 준비"
    (r'^(\d+)\.\s+([가-힣A-Za-z].+)', 'section'),                # "1. 목적" (점 있음)
    
    # 🔥 숫자형 (점 없음): "1 목적", "5 절차" 형식
    (r'^(\d+)\s+([가-힣A-Za-z].+)', 'section'),                  # "1 목적" (공백 1개 이상)
    
    # 🔥 숫자 없는 주요 섹션 (이 문서 형식)
    (r'^(목적)\s*(Purpose)?', 'named_section'),
    (r'^(적용\s*범위)\s*(Scope)?', 'named_section'),
    (r'^(정의)\s*(Definitions)?', 'named_section'),
    (r'^(책임)\s*(Responsibilities)?', 'named_section'),
    (r'^(절차)\s*(Procedure)?', 'named_section'),
    (r'^(기타)\s*(.+)?', 'named_section'),
]


def _is_article_document(text: str) -> bool:
    """조항 기반 문서인지 감지"""
    patterns = [
        r'제\s*\d+\s*조',
        r'제\s*\d+\s*장',
        r'제\s*\d+\s*절',
        r'^\d+\.\d+\.\d+',  # 5.1.1 형식
        r'^\d+\.\d+',       # 5.1 형식
        r'^SOP[-_]?\d+',
    ]

    count = 0
    for pattern in patterns:
        matches = re.findall(pattern, text, re.MULTILINE)
        count += len(matches)

    return count >= 3


def _extract_article_blocks(text: str) -> List[ContentBlock]:
    """
    조항 단위 블록 추출 (SOP 경계 감지 + section_path 계층 추적)
    
    🔥 핵심 기능:
    - section_path: "5 > 5.1 > 5.1.1"
    - section_path_readable: "5 절차 > 5.1 문서체계 > 5.1.1 Level 1"
    """
    lines = text.split('\n')
    blocks = []
    current_lines = []
    current_sop_id = ""  # 현재 SOP ID
    current_meta = {"article_num": None, "article_type": "intro", "title": ""}
    
    sop_id_pattern = re.compile(r'(SOP[-_][A-Z]+[-_]\d+)', re.IGNORECASE)

    # 🔥 계층 추적용 스택
    section_stack = {
        "section": {"num": None, "title": ""},           # 5
        "subsection": {"num": None, "title": ""},        # 5.1
        "subsubsection": {"num": None, "title": ""},     # 5.1.1
    }
    
    # 한글 조항용 스택 (제N조, 제N장 등)
    korean_stack = {
        "chapter": {"num": None, "title": ""},   # 제N장
        "article": {"num": None, "title": ""},   # 제N조
    }

    def build_section_path() -> Dict[str, str]:
        """현재 스택에서 section_path 생성"""
        path_parts = []
        path_readable_parts = []
        
        # 숫자 기반 (5.1.1 형식)
        if section_stack["section"]["num"]:
            path_parts.append(section_stack["section"]["num"])
            title = section_stack["section"]["title"]
            path_readable_parts.append(f"{section_stack['section']['num']} {title}" if title else section_stack["section"]["num"])
        
        if section_stack["subsection"]["num"]:
            path_parts.append(section_stack["subsection"]["num"])
            title = section_stack["subsection"]["title"]
            path_readable_parts.append(f"{section_stack['subsection']['num']} {title}" if title else section_stack["subsection"]["num"])
        
        if section_stack["subsubsection"]["num"]:
            path_parts.append(section_stack["subsubsection"]["num"])
            title = section_stack["subsubsection"]["title"]
            path_readable_parts.append(f"{section_stack['subsubsection']['num']} {title}" if title else section_stack["subsubsection"]["num"])
        
        # 한글 조항 기반 (제N장 > 제N조)
        if korean_stack["chapter"]["num"]:
            ch_num = korean_stack["chapter"]["num"]
            ch_title = korean_stack["chapter"]["title"]
            path_parts.append(f"제{ch_num}장")
            path_readable_parts.append(f"제{ch_num}장 {ch_title}" if ch_title else f"제{ch_num}장")
        
        if korean_stack["article"]["num"]:
            art_num = korean_stack["article"]["num"]
            art_title = korean_stack["article"]["title"]
            path_parts.append(f"제{art_num}조")
            path_readable_parts.append(f"제{art_num}조 {art_title}" if art_title else f"제{art_num}조")
        
        return {
            "section_path": " > ".join(path_parts) if path_parts else None,
            "section_path_readable": " > ".join(path_readable_parts) if path_readable_parts else None,
        }

    def flush():
        if current_lines:
            block_text = '\n'.join(current_lines).strip()
            if block_text:
                # section_path 정보 추가
                path_info = build_section_path()
                
                blocks.append(ContentBlock(
                    text=block_text,
                    block_type=current_meta["article_type"],
                    section=current_meta["article_num"],
                    metadata={
                        "article_num": current_meta["article_num"],
                        "article_type": current_meta["article_type"],
                        "title": current_meta.get("title", ""),
                        "sop_id": current_sop_id,
                        "section_path": path_info["section_path"],
                        "section_path_readable": path_info["section_path_readable"],
                    }
                ))

    for line in lines:
        line_strip = line.strip()
        
        # SOP ID 추출 - 새 SOP 시작이면 현재 블록 flush 먼저!
        sop_match = sop_id_pattern.search(line_strip)
        if sop_match:
            new_sop_id = sop_match.group(1).upper().replace('_', '-')
            if new_sop_id != current_sop_id:
                # 새 SOP 시작 → 현재 블록 저장 후 SOP ID 갱신
                flush()
                current_lines = []
                current_meta = {"article_num": None, "article_type": "intro", "title": ""}
                current_sop_id = new_sop_id
                # 스택 초기화
                section_stack = {
                    "section": {"num": None, "title": ""},
                    "subsection": {"num": None, "title": ""},
                    "subsubsection": {"num": None, "title": ""},
                }
                korean_stack = {
                    "chapter": {"num": None, "title": ""},
                    "article": {"num": None, "title": ""},
                }
        
        # 조항 패턴 매칭
        matched = False
        for pattern, a_type in ARTICLE_PATTERNS:
            m = re.match(pattern, line_strip)
            if m:
                flush()
                current_lines = [line]
                
                num = m.group(1)
                title = m.group(2).strip() if m.lastindex and m.lastindex >= 2 else ""
                
                # 🔥 스택 업데이트
                if a_type == "section":        # 5. 절차
                    section_stack["section"] = {"num": num, "title": title}
                    section_stack["subsection"] = {"num": None, "title": ""}
                    section_stack["subsubsection"] = {"num": None, "title": ""}
                
                elif a_type == "subsection":   # 5.1 문서체계
                    section_stack["subsection"] = {"num": num, "title": title}
                    section_stack["subsubsection"] = {"num": None, "title": ""}
                
                elif a_type == "subsubsection":  # 5.1.1 Level 1
                    section_stack["subsubsection"] = {"num": num, "title": title}
                
                elif a_type == "chapter":      # 제N장
                    korean_stack["chapter"] = {"num": num, "title": title}
                    korean_stack["article"] = {"num": None, "title": ""}
                
                elif a_type == "article":      # 제N조
                    korean_stack["article"] = {"num": num, "title": title}
                
                # 🔥 새 패턴 처리
                elif a_type == "level":        # 제 N레벨
                    section_stack["subsection"] = {"num": f"Level {num}", "title": title}
                    section_stack["subsubsection"] = {"num": None, "title": ""}
                
                elif a_type == "named_section":  # 목적, 적용범위 등
                    # 주요 섹션으로 처리 (스택 리셋)
                    section_stack["section"] = {"num": num, "title": title}
                    section_stack["subsection"] = {"num": None, "title": ""}
                    section_stack["subsubsection"] = {"num": None, "title": ""}
                
                current_meta = {
                    "article_num": num,
                    "article_type": a_type,
                    "title": title
                }
                matched = True
                break

        if not matched:
            current_lines.append(line)

    flush()
    return blocks


def _parse_articles(text: str, filename: str, ext: str) -> ParsedDocument:
    """조항 단위 파싱"""
    blocks = _extract_article_blocks(text)
    metadata = {
        "file_name": filename,
        "file_type": ext,
        "title": _extract_title(text),
        "parser": "article"
    }
    metadata.update(_extract_sop_metadata(text))

    return ParsedDocument(text=text, blocks=blocks, metadata=metadata)


def _parse_plain_text(text: str, filename: str, ext: str) -> ParsedDocument:
    """단순 텍스트 파싱"""
    paragraphs = re.split(r'\n\s*\n', text)
    blocks = []

    for p in paragraphs:
        p = p.strip()
        if p:
            blocks.append(ContentBlock(text=p, block_type="paragraph"))

    return ParsedDocument(
        text=text,
        blocks=blocks,
        metadata={
            "file_name": filename,
            "file_type": ext,
            "title": _extract_title(text),
            "parser": "plain"
        }
    )


# ═══════════════════════════════════════════════════════════════════════════
# 메타데이터 추출
# ═══════════════════════════════════════════════════════════════════════════

def _extract_title(text: str) -> str:
    """문서 제목 추출"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:10]:
        if line.lower().startswith(("title:", "제목:")):
            return line.split(':', 1)[1].strip()
        if re.match(r'^SOP[-_]?\d+', line, re.IGNORECASE):
            return line[:100]
        if len(line) > 5 and not line.startswith('#'):
            return line[:100]
    return "제목 없음"


def _extract_sop_metadata(text: str) -> Dict:
    """SOP 관련 메타데이터 추출"""
    metadata = {}

    # SOP ID
    sop_match = re.search(r'(SOP[-_]?[A-Z]*[-_]?\d+)', text, re.IGNORECASE)
    if sop_match:
        metadata["sop_id"] = sop_match.group(1)

    # 버전
    ver_match = re.search(r'(?:Version|Ver|버전|개정)[\s.:]*(\d+\.?\d*)', text, re.IGNORECASE)
    if ver_match:
        metadata["version"] = ver_match.group(1)

    # 부서
    dept_match = re.search(r'(?:부서|Dept|Department)[\s:]*([가-힣\w\s]+?)(?:\n|$)', text, re.IGNORECASE)
    if dept_match:
        metadata["dept"] = dept_match.group(1).strip()

    # 시행일
    date_match = re.search(r'(?:시행일|Effective|발효)[\s:]*(\d{4}[-./]\d{1,2}[-./]\d{1,2})', text, re.IGNORECASE)
    if date_match:
        metadata["effective_date"] = date_match.group(1)

    return metadata


# ═══════════════════════════════════════════════════════════════════════════
# Fallback 파서들 (Docling 없을 때)
# ═══════════════════════════════════════════════════════════════════════════

def _decode_text(content: bytes) -> str:
    """바이트 → 텍스트 디코딩"""
    for encoding in ["utf-8", "cp949", "euc-kr", "latin-1"]:
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="ignore")


def _load_pdf_basic(filename: str, content: bytes) -> ParsedDocument:
    """기본 PDF 파싱 (PyMuPDF + OCR fallback)"""
    try:
        import fitz
    except ImportError:
        return ParsedDocument(
            text="PDF 파싱 라이브러리(PyMuPDF)가 설치되지 않았습니다.",
            blocks=[],
            metadata={"file_name": filename, "error": "pymupdf not installed"}
        )

    doc = fitz.open(stream=content, filetype="pdf")
    blocks = []
    all_text = []
    tables = []

    for page_idx, page in enumerate(doc):
        text = page.get_text().strip()

        # 텍스트 없으면 OCR 실행
        if not text:
            ocr_text = _ocr_pdf_page(page)
            if ocr_text:
                text = ocr_text
                source = "pymupdf+ocr"
            else:
                continue
        else:
            source = "pymupdf"

        blocks.append(ContentBlock(
            text=text,
            block_type="page",
            page=page_idx + 1,
            metadata={"source": source}
        ))
        all_text.append(text)

        # 표 추출 시도
        try:
            page_tables = page.find_tables()
            for table in page_tables:
                tables.append({
                    "page": page_idx + 1,
                    "rows": table.extract()
                })
        except Exception:
            pass

    full_text = "\n\n".join(all_text)

    # 조항 재파싱
    if _is_article_document(full_text):
        article_blocks = _extract_article_blocks(full_text)
        if article_blocks:
            blocks = article_blocks

    metadata = {
        "file_name": filename,
        "file_type": "pdf",
        "total_pages": len(doc),
        "title": _extract_title(full_text),
        "table_count": len(tables),
        "parser": "pymupdf+ocr"
    }
    metadata.update(_extract_sop_metadata(full_text))

    return ParsedDocument(
        text=full_text,
        blocks=blocks,
        metadata=metadata,
        tables=tables
    )


def _load_html_basic(filename: str, content: bytes) -> ParsedDocument:
    """기본 HTML 파싱"""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        text = _decode_text(content)
        return _parse_plain_text(text, filename, ".html")

    html = _decode_text(content)
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style"]):
        tag.decompose()

    blocks = []
    tables = []

    # 표 추출
    for table in soup.find_all('table'):
        rows = []
        for tr in table.find_all('tr'):
            cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
            rows.append(cells)
        if rows:
            tables.append({"rows": rows})
            blocks.append(ContentBlock(
                text='\n'.join([' | '.join(r) for r in rows]),
                block_type="table"
            ))
        table.decompose()

    # 나머지 텍스트
    text = soup.get_text('\n', strip=True)
    for para in text.split('\n\n'):
        para = para.strip()
        if para:
            blocks.append(ContentBlock(text=para, block_type="paragraph"))

    full_text = soup.get_text('\n', strip=True)

    return ParsedDocument(
        text=full_text,
        blocks=blocks,
        metadata={
            "file_name": filename,
            "file_type": "html",
            "title": _extract_title(full_text),
            "table_count": len(tables),
            "parser": "beautifulsoup"
        },
        tables=tables
    )


def _load_csv(filename: str, content: bytes) -> ParsedDocument:
    """CSV 파싱"""
    text = _decode_text(content)
    lines = text.splitlines()

    rows = []
    for line in lines:
        cells = [c.strip('" ') for c in line.split(',')]
        rows.append(cells)

    table_text = '\n'.join([' | '.join(row) for row in rows])

    return ParsedDocument(
        text=table_text,
        blocks=[ContentBlock(text=table_text, block_type="table")],
        metadata={
            "file_name": filename,
            "file_type": "csv",
            "title": filename,
            "row_count": len(rows),
            "parser": "csv"
        },
        tables=[{"rows": rows}]
    )


def _ocr_pdf_page(page, ocr=None, dpi=300):
    """
    PyMuPDF page → OCR 텍스트
    ocr: RapidOCR 객체를 외부에서 전달 가능 (재사용 권장)
    """
    import tempfile
    import os
    
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        print("⚠️ RapidOCR not installed, skipping OCR")
        return ""

    if ocr is None:
        ocr = RapidOCR()

    pix = page.get_pixmap(dpi=dpi)

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        pix.save(tmp.name)
        img_path = tmp.name

    try:
        result, _ = ocr(img_path)
        if result:
            texts = [text for _, text, score in result if score > 0.5]
            return "\n".join(texts)
        return ""
    except Exception as e:
        print(f"⚠️ OCR failed: {e}")
        return ""
    finally:
        if os.path.exists(img_path):
            os.remove(img_path)